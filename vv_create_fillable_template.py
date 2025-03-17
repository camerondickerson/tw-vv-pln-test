from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def create_fillable_template():
    # Create a new Document
    doc = Document()

    # Set the styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title
    title = doc.add_heading('Software Verification Plan and Protocol', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Document Owner
    doc.add_paragraph('Document Owner: Software Test')

    # Document Identification
    doc.add_heading('DOCUMENT IDENTIFICATION:', level=2)
    doc.add_paragraph('Software Verification Plan and Protocol')

    # Instructions
    doc.add_heading('Instructions', level=2)
    doc.add_paragraph('Title Guidelines\n'
                      'If the document is for the Software Verification Plan and Protocol, use below as an example:\n'
                      'SW##### [name of the software] Software Verification Plan and Protocol\n'
                      'If the document is for the Software Compatibility Verification Plan and Protocol, use below as an example:\n'
                      'SW##### [name of the software] Software Compatibility Verification Plan and Protocol for Version A.B.C.D with [new device OS version, new device model, etc.]\n'
                      'If the document is for the System End to End Verification Plan and Protocol, use below as an example:\n'
                      '[name of the product/project] Software System End to End Verification Plan and Protocol')

    # Purpose
    doc.add_heading('PURPOSE', level=2)
    doc.add_paragraph('[Describe the purpose of this document.]')

    # System Overview
    doc.add_heading('SYSTEM OVERVIEW', level=2)
    doc.add_paragraph('[Provide the system overview of the project.]')

    # Definitions and Acronyms
    doc.add_heading('DEFINITIONS AND ACRONYMS', level=2)
    doc.add_paragraph('Term: Definition\n'
                      'Functional Testing: Functional testing involves testing the application against the business requirements.')

    # References
    doc.add_heading('REFERENCES', level=2)
    doc.add_paragraph('Document No.: Description: Rev\n'
                      'IEC 62304+AMD1: Medical Device Software - Software life cycle processes: N/A')

    # Scope
    doc.add_heading('SCOPE', level=2)
    doc.add_paragraph('[Describe the new features in scope for test. Identify new or updated requirements to be tested. Include the Enhancement and/or Anomaly ID (if applicable) and the description of the change.]')

    # Tools and Sample Size
    doc.add_heading('TOOLS AND SAMPLE SIZE', level=2)
    doc.add_paragraph('[List the software and hardware tools necessary for testing (e.g., hardware fixtures, hardware accessories, third party and/or in-house tools). Include Software Configuration Management (SCM) tools used and version control location/information.]')

    # Test Methodology
    doc.add_heading('TEST METHODOLOGY', level=2)
    doc.add_paragraph('[At a minimum each Software Requirement Specification (SRS) requirement must have at least one corresponding Test Case that demonstrates that the requirement has been met. This correspondence (requirement to test and test to test result) must be confirmed with the Requirements Trace Matrix.]')

    # Anomaly Tracking
    doc.add_heading('ANOMALY TRACKING', level=2)
    doc.add_paragraph('[Describe the strategy for anomaly management.]')

    # Acceptance Criteria
    doc.add_heading('ACCEPTANCE CRITERIA', level=2)
    doc.add_paragraph('[Describe the acceptance criteria for the software release.]')

    # Attachments
    doc.add_heading('ATTACHMENTS', level=2)
    doc.add_paragraph('[List attachments to this document (i.e., test cases/scripts with links to the requirements and/or anomalies and test case/script peer review record).]')

    # Save the document
    directory = os.path.expanduser('~/Documents')
    if not os.path.exists(directory):
        os.makedirs(directory)
    doc.save(os.path.join(directory, 'Software_Verification_Plan_and_Protocol_Template.docx'))

if __name__ == "__main__":
    create_fillable_template()